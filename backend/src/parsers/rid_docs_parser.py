from docx import Document
import re

def analyze_template(template_path):
    """
    Анализ шаблона документа и поиск мест для замены, включая поля для сторон соглашения
    """
    doc = Document(template_path)
    placeholders = []

    # Шаблоны для поиска мест заполнения
    patterns = [
        r'[«"]?_+[»"]?',  # Подчеркивания
        r'именуем\w{1,3}\s+(?:в\s+дальнейшем\s+)?[«"]([^»"]+)[»"]',  # Именования сторон
        r'лице\s+([^,]+),',  # Представители в лице
        r'основании\s+([^,\.]+)',  # Основание полномочий
        r'местонахожден\w+:\s*([^,\.]+)',  # Адреса
        r'(?:ОГРН|ИНН|КПП)\s*:?\s*(\d*_*\d*)',  # Реквизиты
        r'№\s*[\d_]*\s*от\s*[«"]?[\d\._]*[»"]?',  # Номера и даты документов
        r'[«"]([^»"]+)[»"]\s*\(далее\s*[-—]\s*[^)]+\)',  # Определения терминов
        r'стоимост\w+\s+[\d\s_,]+\s*(?:руб\.|₽)?',  # Стоимость
        r'\d{2}[\.]\d{2}[\.]\d{4}',  # Даты в формате ДД.ММ.ГГГГ
    ]

    for paragraph in doc.paragraphs:
        for pattern in patterns:
            # Поиск по каждому паттерну
            matches = re.finditer(pattern, paragraph.text)
            for match in matches:
                # Сохраняем информацию о форматировании
                start, end = match.span()
                formatted_runs = []
                current_position = 0
                
                for run in paragraph.runs:
                    run_length = len(run.text)
                    if current_position + run_length > start and current_position < end:
                        formatted_runs.append({
                            'text': run.text,
                            'bold': run.bold,
                            'italic': run.italic,
                            'underline': run.underline,
                            'position': (current_position, current_position + run_length)
                        })
                    current_position += run_length

                placeholders.append({
                    'text': paragraph.text,
                    'match': match.group(0),
                    'pattern_type': pattern,
                    'formatted_runs': formatted_runs
                })

    return placeholders

def replace_placeholders(template_path, output_path, replacement_data):
    """
    Замена найденных плейсхолдеров с сохранением форматирования
    """
    doc = Document(template_path)

    # Замена в параграфах с сохранением форматирования
    for paragraph in doc.paragraphs:
        runs = paragraph.runs
        for i, run in enumerate(runs):
            for key, value in replacement_data.items():
                if key in run.text:
                    # Сохраняем форматирование
                    bold = run.bold
                    italic = run.italic
                    underline = run.underline
                    font_name = run.font.name
                    font_size = run.font.size
                    
                    # Заменяем текст
                    run.text = run.text.replace(key, value)
                    
                    # Восстанавливаем форматирование
                    run.bold = bold
                    run.italic = italic
                    run.underline = underline
                    run.font.name = font_name
                    run.font.size = font_size

    # Замена в таблицах с сохранением форматирования
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    runs = paragraph.runs
                    for i, run in enumerate(runs):
                        for key, value in replacement_data.items():
                            if key in run.text:
                                # Сохраняем форматирование
                                bold = run.bold
                                italic = run.italic
                                underline = run.underline
                                font_name = run.font.name
                                font_size = run.font.size
                                
                                # Заменяем текст
                                run.text = run.text.replace(key, value)
                                
                                # Восстанавливаем форматирование
                                run.bold = bold
                                run.italic = italic
                                run.underline = underline
                                run.font.name = font_name
                                run.font.size = font_size

    doc.save(output_path)

def main():
    template_path = "/home/dukhanin/fic/docs/Приказ_и_иные_документы_на_создание_РИД_1 .docx"
    output_path = "/home/dukhanin/fic/docs/Приказ_РИД_заполненный.docx"

    # Анализ шаблона
    placeholders = analyze_template(template_path)
    
    # Вывод найденных мест для замены
    print("Найдены следующие места для заполнения:")
    for idx, placeholder in enumerate(placeholders, 1):
        print(f"{idx}. {placeholder['text']}")
        print(f"   Найдено: {placeholder['match']}")
        if placeholder['formatted_runs']:
            print("   Форматирование:")
            for run in placeholder['formatted_runs']:
                print(f"      - {run['text']}: bold={run['bold']}, italic={run['italic']}, underline={run['underline']}")
        print()

    # Пример данных для замены (нужно адаптировать под конкретный документ)
    replacement_data = {
        '___': 'ООО "Инновационные Технологии"',
        '_____': 'Петров Петр Петрович',
        '__.__.____': '01.01.2024',
        '№___': '№123',
        '_________': 'Разработка системы ИИ',
    }

    # Замена и сохранение
    replace_placeholders(template_path, output_path, replacement_data)
    print(f"Документ сохранен: {output_path}")

if __name__ == "__main__":
    main() 